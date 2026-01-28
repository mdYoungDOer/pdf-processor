"""
PDF Table Extraction Processor Module

This module provides stateless PDF table extraction functionality.
All operations are performed in-memory without disk I/O.
"""

import io
import re
import pandas as pd
import pdfplumber


def _make_unique_columns(columns):
    """Make column names unique by adding suffixes to duplicates."""
    seen = {}
    result = []
    for col in columns:
        # Convert None to empty string
        col = str(col).strip() if col is not None else ''
        if col in seen:
            seen[col] += 1
            result.append(f"{col}_{seen[col]}")
        else:
            seen[col] = 0
            result.append(col)
    return result


def _count_non_empty_cells(row):
    """Count non-empty cells in a row."""
    count = 0
    for cell in row:
        if cell is not None and str(cell).strip():
            count += 1
    return count


def _get_non_empty_indices(row):
    """Get indices of non-empty cells in a row."""
    indices = []
    for i, cell in enumerate(row):
        if cell is not None and str(cell).strip():
            indices.append(i)
    return indices


def _find_table_structure(table):
    """
    Analyze table to find the dominant column structure.
    
    Returns:
        Tuple of (header_row_index, columns_to_keep)
    """
    if not table or len(table) < 2:
        return 0, None
    
    # Analyze each row to find patterns
    row_profiles = []
    for i, row in enumerate(table):
        non_empty = _count_non_empty_cells(row)
        indices = _get_non_empty_indices(row)
        row_profiles.append({
            'index': i,
            'non_empty': non_empty,
            'indices': indices,
            'total_cols': len(row)
        })
    
    # Find the most common number of non-empty cells (excluding very sparse rows)
    # Lower threshold to 2 to catch more valid rows
    non_empty_counts = [p['non_empty'] for p in row_profiles if p['non_empty'] >= 2]
    
    if not non_empty_counts:
        # If no rows with 2+ cells, return first row as header and keep all columns
        return 0, None
    
    # Find the mode (most common column count)
    from collections import Counter
    count_freq = Counter(non_empty_counts)
    # Get the most common count that's reasonably high
    if count_freq:
        target_col_count = max(count_freq.keys(), key=lambda x: (count_freq[x], x))
    else:
        target_col_count = 3  # Default fallback
    
    # Find the first row that matches this structure (likely the header)
    # Lower tolerance to 0.5 to be less strict
    header_idx = 0
    for profile in row_profiles:
        if profile['non_empty'] >= max(2, target_col_count * 0.5):  # Lower tolerance
            header_idx = profile['index']
            break
    
    # Analyze data rows (after header) to find which columns are actually used
    data_rows = table[header_idx + 1:] if header_idx + 1 < len(table) else []
    
    if not data_rows:
        return header_idx, None
    
    # Count how often each column position has data
    max_cols = max(len(row) for row in table)
    col_usage = [0] * max_cols
    
    for row in data_rows:
        for col_idx, cell in enumerate(row):
            if cell is not None and str(cell).strip():
                col_usage[col_idx] += 1
    
    # Lower threshold to 15% to be less aggressive
    threshold = max(1, len(data_rows) * 0.15)
    cols_to_keep = [i for i, usage in enumerate(col_usage) if usage >= threshold]
    
    # Also include columns from the header row that have content
    header_row = table[header_idx] if header_idx < len(table) else []
    for i, cell in enumerate(header_row):
        if cell is not None and str(cell).strip() and i not in cols_to_keep:
            # Only add if this column has at least some data
            if i < len(col_usage) and col_usage[i] > 0:
                cols_to_keep.append(i)
    
    # If we filtered out too many columns (less than 3), keep more
    if len(cols_to_keep) < 3 and max_cols > 3:
        # Keep top columns by usage
        col_usage_with_idx = [(usage, i) for i, usage in enumerate(col_usage)]
        col_usage_with_idx.sort(reverse=True)
        cols_to_keep = [i for _, i in col_usage_with_idx[:max(3, max_cols // 2)]]
        # Also include header columns
        for i, cell in enumerate(header_row):
            if cell is not None and str(cell).strip() and i not in cols_to_keep:
                cols_to_keep.append(i)
    
    cols_to_keep = sorted(set(cols_to_keep))
    
    return header_idx, cols_to_keep


def _clean_table(table):
    """
    Clean a table by removing graphical header rows and artifact columns.
    
    Args:
        table: List of rows (each row is a list of cells)
        
    Returns:
        Cleaned table
    """
    if not table:
        return table
    
    # Find the actual table structure
    header_idx, cols_to_keep = _find_table_structure(table)
    
    # Use rows from header_idx onwards
    cleaned_table = table[header_idx:]
    
    if not cleaned_table:
        return table
    
    # If we couldn't determine columns to keep, try a simpler approach
    if cols_to_keep is None or not cols_to_keep:
        # Fall back to keeping columns with at least some data
        num_rows = len(cleaned_table)
        max_cols = max(len(row) for row in cleaned_table) if cleaned_table else 0
        
        if max_cols == 0:
            return cleaned_table
        
        col_counts = [0] * max_cols
        for row in cleaned_table:
            for col_idx, cell in enumerate(row):
                if cell is not None and str(cell).strip():
                    col_counts[col_idx] += 1
        
        # Lower threshold to 20% fill rate
        threshold = max(1, num_rows * 0.2)
        cols_to_keep = [i for i, count in enumerate(col_counts) if count >= threshold]
        
        # If still no columns, keep at least the first few columns that have any data
        if not cols_to_keep:
            cols_to_keep = [i for i, count in enumerate(col_counts) if count > 0]
            # Limit to reasonable number (max 20 columns)
            if len(cols_to_keep) > 20:
                cols_with_counts = [(col_counts[i], i) for i in cols_to_keep]
                cols_with_counts.sort(reverse=True)
                cols_to_keep = [i for _, i in cols_with_counts[:20]]
    
    if not cols_to_keep:
        # Last resort: return original table if we can't determine columns
        return cleaned_table
    
    # Filter to keep only the identified columns
    result = []
    for row in cleaned_table:
        new_row = []
        for col_idx in cols_to_keep:
            if col_idx < len(row):
                new_row.append(row[col_idx])
            else:
                new_row.append(None)
        result.append(new_row)
    
    return result


def tables_to_dataframe(file_bytes: bytes) -> pd.DataFrame:
    """
    Extract all tables from a PDF file and merge them into a single DataFrame.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        pd.DataFrame: A single merged DataFrame containing all tables from all pages.
                     Returns an empty DataFrame if no tables are found.
        
    Raises:
        ValueError: If the PDF file is empty or invalid
        Exception: If the PDF file is corrupted or cannot be processed
    """
    if not file_bytes:
        raise ValueError("PDF file bytes cannot be empty")
    
    if len(file_bytes) == 0:
        raise ValueError("PDF file is empty")
    
    # Create BytesIO wrapper for in-memory PDF processing
    pdf_stream = io.BytesIO(file_bytes)
    
    all_tables = []
    
    try:
        with pdfplumber.open(pdf_stream) as pdf:
            # Iterate through all pages
            for page in pdf.pages:
                # Extract tables from current page with table settings
                # Use more precise table detection
                tables = page.extract_tables(table_settings={
                    "vertical_strategy": "lines",
                    "horizontal_strategy": "lines",
                    "snap_tolerance": 3,
                    "join_tolerance": 3,
                })
                
                # If no tables found with strict settings, try with text strategy
                if not tables:
                    tables = page.extract_tables(table_settings={
                        "vertical_strategy": "text",
                        "horizontal_strategy": "text",
                    })
                
                if tables:
                    # Convert each table to a DataFrame
                    for table in tables:
                        if table and len(table) > 0:
                            try:
                                # Clean the table (remove graphical header artifacts)
                                cleaned = _clean_table(table)
                                
                                # Ensure we have a valid cleaned table
                                if not cleaned or len(cleaned) == 0:
                                    # Fallback to original table if cleaning removed everything
                                    cleaned = table
                                
                                if cleaned and len(cleaned) > 1:
                                    # First row as headers
                                    columns = _make_unique_columns(cleaned[0])
                                    df = pd.DataFrame(cleaned[1:], columns=columns)
                                    
                                    # Remove rows that are completely empty
                                    df = df.dropna(how='all')
                                    
                                    # Remove columns that are completely empty
                                    df = df.dropna(axis=1, how='all')
                                    
                                    # Only add if we have at least some data
                                    if not df.empty and len(df.columns) > 0:
                                        all_tables.append(df)
                                elif cleaned and len(cleaned) == 1:
                                    # Only header row
                                    columns = _make_unique_columns(cleaned[0])
                                    # Filter out empty column names
                                    columns = [c for c in columns if c and str(c).strip()]
                                    if len(columns) > 0:
                                        df = pd.DataFrame(columns=columns)
                                        all_tables.append(df)
                            except Exception as e:
                                # If cleaning fails, try to use original table
                                try:
                                    if len(table) > 1:
                                        columns = _make_unique_columns(table[0])
                                        df = pd.DataFrame(table[1:], columns=columns)
                                        df = df.dropna(how='all')
                                        df = df.dropna(axis=1, how='all')
                                        if not df.empty:
                                            all_tables.append(df)
                                except:
                                    # Skip this table if we can't process it
                                    pass
    
    except Exception as e:
        raise Exception(f"Error processing PDF: {str(e)}")
    
    # Merge all tables into a single DataFrame
    if all_tables:
        # Use axis=0 concatenation with join='outer' to handle different columns
        merged_df = pd.concat(all_tables, ignore_index=True, axis=0, join='outer')
        
        # Final cleanup: remove columns that are mostly empty
        # Lower threshold to 20% to be less aggressive
        threshold = len(merged_df) * 0.2
        merged_df = merged_df.dropna(axis=1, thresh=max(1, int(threshold)))
        
        # Also remove columns that have only empty strings
        for col in merged_df.columns:
            if merged_df[col].apply(lambda x: str(x).strip() if pd.notna(x) else '').eq('').all():
                merged_df = merged_df.drop(columns=[col])
        
        # Remove columns with empty or whitespace-only names if they have little data
        cols_to_drop = []
        for col in merged_df.columns:
            col_str = str(col).strip()
            if col_str == '' or col_str == 'None':
                # Check if this column has meaningful data
                non_empty = merged_df[col].apply(lambda x: bool(str(x).strip()) if pd.notna(x) else False).sum()
                if non_empty < len(merged_df) * 0.3:
                    cols_to_drop.append(col)
        
        if cols_to_drop:
            merged_df = merged_df.drop(columns=cols_to_drop)
        
        return merged_df
    else:
        # Return empty DataFrame if no tables found
        return pd.DataFrame()


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """
    Extract all text from a PDF file and format it nicely.
    
    Args:
        file_bytes: PDF file content as bytes
        
    Returns:
        str: Formatted text content from all pages
        
    Raises:
        ValueError: If the PDF file is empty or invalid
        Exception: If the PDF file is corrupted or cannot be processed
    """
    if not file_bytes:
        raise ValueError("PDF file bytes cannot be empty")
    
    if len(file_bytes) == 0:
        raise ValueError("PDF file is empty")
    
    # Create BytesIO wrapper for in-memory PDF processing
    pdf_stream = io.BytesIO(file_bytes)
    
    all_text = []
    
    try:
        with pdfplumber.open(pdf_stream) as pdf:
            # Iterate through all pages
            for page_num, page in enumerate(pdf.pages, 1):
                # Extract text from current page
                page_text = page.extract_text()
                
                if page_text:
                    # Clean and format the text
                    cleaned_text = _format_text(page_text, page_num, len(pdf.pages))
                    all_text.append(cleaned_text)
    
    except Exception as e:
        raise Exception(f"Error processing PDF: {str(e)}")
    
    return "\n\n".join(all_text)


def _format_text(text: str, page_num: int, total_pages: int) -> str:
    """
    Format extracted text nicely.
    
    Args:
        text: Raw text from PDF page
        page_num: Current page number
        total_pages: Total number of pages
        
    Returns:
        Formatted text string
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    text = re.sub(r'\n{3,}', '\n\n', text)  # Max 2 consecutive newlines
    text = re.sub(r'[ \t]+', ' ', text)  # Normalize spaces
    
    # Add page separator if multiple pages
    if total_pages > 1:
        header = f"\n{'='*60}\nPage {page_num} of {total_pages}\n{'='*60}\n\n"
        text = header + text
    
    # Clean up line breaks
    lines = text.split('\n')
    formatted_lines = []
    
    for line in lines:
        line = line.strip()
        if line:
            formatted_lines.append(line)
        elif formatted_lines and formatted_lines[-1]:  # Add blank line only if previous line wasn't blank
            formatted_lines.append('')
    
    return '\n'.join(formatted_lines)


def convert_to_format(df: pd.DataFrame, output_format: str) -> io.BytesIO:
    """
    Convert a Pandas DataFrame to the specified format (CSV or Excel).
    
    Args:
        df: Pandas DataFrame to convert
        output_format: Output format, either 'csv' or 'excel'
        
    Returns:
        io.BytesIO: BytesIO stream containing the formatted data
        
    Raises:
        ValueError: If output_format is not 'csv' or 'excel'
        Exception: If conversion fails
    """
    if output_format not in ['csv', 'excel']:
        raise ValueError(f"Invalid output_format: {output_format}. Must be 'csv' or 'excel'")
    
    if df.empty:
        # Return empty BytesIO for empty DataFrames
        return io.BytesIO()
    
    output_stream = io.BytesIO()
    
    try:
        if output_format == 'csv':
            # Convert DataFrame to CSV
            csv_string = df.to_csv(index=False)
            output_stream.write(csv_string.encode('utf-8'))
            output_stream.seek(0)  # Reset stream position to beginning
            
        elif output_format == 'excel':
            # Convert DataFrame to Excel using openpyxl engine
            with pd.ExcelWriter(output_stream, engine='openpyxl') as writer:
                df.to_excel(writer, index=False, sheet_name='Sheet1')
            output_stream.seek(0)  # Reset stream position to beginning
            
    except Exception as e:
        raise Exception(f"Error converting DataFrame to {output_format}: {str(e)}")
    
    return output_stream


def convert_text_to_format(text: str, output_format: str) -> io.BytesIO:
    """
    Convert extracted text to the specified format (TXT or DOCX).
    
    Args:
        text: Text content to convert
        output_format: Output format, either 'txt' or 'docx'
        
    Returns:
        io.BytesIO: BytesIO stream containing the formatted data
        
    Raises:
        ValueError: If output_format is not 'txt' or 'docx'
        Exception: If conversion fails
    """
    if output_format not in ['txt', 'docx']:
        raise ValueError(f"Invalid output_format: {output_format}. Must be 'txt' or 'docx'")
    
    if not text or not text.strip():
        # Return empty BytesIO for empty text
        return io.BytesIO()
    
    output_stream = io.BytesIO()
    
    try:
        if output_format == 'txt':
            # Convert text to plain text file
            output_stream.write(text.encode('utf-8'))
            output_stream.seek(0)  # Reset stream position to beginning
            
        elif output_format == 'docx':
            # Convert text to DOCX using python-docx
            from docx import Document
            from docx.shared import Pt
            
            doc = Document()
            
            # Set default font
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Calibri'
            font.size = Pt(11)
            
            # Split text into paragraphs and add to document
            paragraphs = text.split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    # Handle page headers
                    if para_text.strip().startswith('='):
                        p = doc.add_paragraph(para_text.strip())
                        p.style = 'Heading 1'
                    else:
                        # Split by single newlines within paragraph
                        lines = para_text.split('\n')
                        for i, line in enumerate(lines):
                            if line.strip():
                                p = doc.add_paragraph(line.strip())
                                if i == 0 and len(lines) > 1:
                                    # First line might be a heading
                                    if len(line.strip()) < 100 and not line.strip().endswith('.'):
                                        p.style = 'Heading 2'
            
            # Save to BytesIO
            doc.save(output_stream)
            output_stream.seek(0)  # Reset stream position to beginning
            
    except Exception as e:
        raise Exception(f"Error converting text to {output_format}: {str(e)}")
    
    return output_stream
